import re
from docx import Document
import pdfplumber
import logging

logging.getLogger("pdfminer").setLevel(logging.ERROR)

PDF_FILTERS = [
    r"#|\$|\*|\(|\)|\[|\]|%|\.|\,|/\d\d\d",
    r"COURSE NAME ISE ESE TOTAL TW PR OR TUT Tot Crd Grd GP CP PR ORD"
]
PDF_SUBSTITUTIONS = {
    r'\sAB\s': ' -- ',
    r'\sAB\s': ' -- ',
    r'\s--\s': ' --- ',
    r'\s--\s': ' --- '
}
PDF_PATTERNS = {
    'name': r"NAME\s:\s([A-Z]+\s[A-Z]+\s[A-Z]+)",
    'course': r"^(\d+[A-Z]?)\s+([A-Z: &]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z+]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)\s+(\d+|---|[A-Z]+)"
}

SCORE_TYPES = ['ISE', 'ESE', 'TOTAL', 'TW', 'PR', 'OR', 'TUT', 'Tot', 'Crd', 'Grd', 'GP', 'CP', 'PAR', 'ORD']

def pdf_parser(file_path, filters, patterns, substitutions, add_callback, progress_callback=None):
    try:
        file = pdfplumber.open(file_path)
    except Exception as e:
        logging.error(e)
        return
    
    for page_no, page in enumerate(file.pages, start=0):
        data = page.extract_text()

        for regex_filter in filters:
            data = re.sub(regex_filter, '', data)
            
        for sub, val in substitutions.items():
            data = re.sub(sub, val, data)
        
        lines = data.splitlines()
        lines_count = len(lines)
        for line_no, line in enumerate(lines, start=1):
            for key, pattern in patterns.items():
                if _match := re.search(pattern, line):
                        if len(_match.groups()) == 1:
                            add_callback(key, _match.group(1))
                        else:
                            add_callback(key, tuple(val.strip() for val in _match.groups()))
            if progress_callback:
                progress_callback(int((page_no * line_no) / (len(file.pages) * lines_count) * 100))

    file.close()

def export_to_docx(file_path, result, options, progress_callback=None):
    doc = Document()
    table = doc.add_table(rows=len(result.keys()) + 1, cols=len(options) + 1)

    table.cell(0, 0).text = 'NAME'
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    for i, (course, score_type) in enumerate(options, start=1):
        table.cell(0, i).text = f'{course} ({score_type})'
        table.cell(0, i).paragraphs[0].runs[0].bold = True

    progress_tot = len(result.items()) - 1
    for i, (name, courses) in enumerate(result.items(), start=1):
        table.cell(i, 0).text = name
        for j, (course, score_type) in enumerate(options, start=1):
            val = courses.get(course, None)
            if val:
                table.cell(i, j).text = val[score_type]
        
        if progress_callback:
            progress_callback((i / progress_tot) * 100)
    
    doc.save(file_path)
